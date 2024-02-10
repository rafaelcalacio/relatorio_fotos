from docx import Document
from docx.shared import Inches
from PIL import Image
import os

# Função para adicionar seis fotos por página
def add_six_pictures(doc, image_paths, width, height):
    num_images = len(image_paths)
    num_pages = (num_images + 5) // 6  # Calcula o número de páginas necessárias

    for page_num in range(num_pages):
        if page_num != 0:
            doc.add_page_break()  # Adiciona uma nova página, exceto na primeira
            
         # Adiciona um parágrafo no início de cada página
        doc.add_paragraph("")
        doc.add_paragraph("Ambiente")

        table = doc.add_table(rows=3, cols=2)  # Cria uma tabela 3x2 para as imagens
        table.autofit = False  # Desativa o ajuste automático da tabela

        for row in table.rows:
            for cell in row.cells:
                if image_paths:
                    image_path = image_paths.pop(0)
                    # Abre a imagem com PIL para verificar a orientação
                    img = Image.open(image_path)
                    # Verifica a orientação e gira a imagem se necessário
                    if hasattr(img, '_getexif'):
                        exif = img._getexif()
                        if exif is not None and 274 in exif:
                            orientation = exif[274]
                            if orientation == 3:
                                img = img.rotate(180, expand=True)
                            elif orientation == 6:
                                img = img.rotate(270, expand=True)
                            elif orientation == 8:
                                img = img.rotate(90, expand=True)
                    img.save(image_path)  # Salva a imagem com a orientação corrigida
                    # Insere a imagem no documento Word
                    run = cell.paragraphs[0].add_run()
                    run.add_picture(image_path, width=width, height=height)

# Crie um novo documento Word
doc = Document()

# Tamanho da página A4 em polegadas
a4_width = Inches(6.5)  # largura
a4_height = Inches(7.5)  # altura

# Calcula a largura e a altura das imagens para preencher a página A4
image_width = a4_width / 2  # 2 colunas
image_height = a4_height / 3  # 3 linhas

# Diretório contendo as imagens
diretorios_imagens =  [f'C:\\Users\\Elenita\\OneDrive\\Pictures\\pics123\\sala',
                      f'C:\\Users\\Elenita\\OneDrive\\Pictures\\pics123\\quarto',
                      f'C:\\Users\\Elenita\\OneDrive\\Pictures\\pics123\\lavabo'] # Exemplo com 3 páginas


# Adicione as fotos ao documento
for diretorio_imagens in diretorios_imagens:
    image_paths = [os.path.join(diretorio_imagens, f"lazer ({i}).jpg") for i in range(1, 7)]  # 6 imagens por página
    add_six_pictures(doc, image_paths, width=image_width, height=image_height)


# Salve o documento
doc.save('relatorio.docx')

print("Imagens inseridas com sucesso no documento Word!")
